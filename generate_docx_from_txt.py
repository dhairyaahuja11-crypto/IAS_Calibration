import docx
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Read the content from the txt file
with open('Comprehensive_Workflow_Overview.docx.txt', 'r', encoding='utf-8') as f:
    lines = f.readlines()

doc = docx.Document()

# Style for headings
heading_map = {
    1: 'Heading 1',
    2: 'Heading 2',
    3: 'Heading 3',
}

def add_heading(text, level=1):
    doc.add_heading(text, level=level)

def add_paragraph(text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

for line in lines:
    line = line.rstrip('\n')
    if line.startswith('# '):
        add_heading(line[2:], 1)
    elif line.startswith('## '):
        add_heading(line[3:], 2)
    elif line.startswith('### '):
        add_heading(line[4:], 3)
    elif line.strip() == '---':
        doc.add_page_break()
    elif line.strip() == '':
        doc.add_paragraph('')
    else:
        add_paragraph(line)

# Save as docx
output_path = 'Comprehensive_Workflow_Overview.docx'
doc.save(output_path)
print(f"Saved as {output_path}")
